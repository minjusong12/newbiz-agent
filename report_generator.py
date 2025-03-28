from datetime import datetime
from docx import Document
from docx.shared import Pt
from research import run_deep_research  # 1단계: GPT 리서치
from generate_slides import generate_slide_outline  # 2단계: 슬라이드 구조
from generate_pptx import create_ppt_from_outline  # 3단계: PPT 생성
import os
import zipfile


# 📄 보고서 텍스트 생성
def generate_business_report(topic: str, research_text: str) -> str:
    today = datetime.today().strftime("%Y-%m-%d")
    report = f"""
[표지]
- 주제: {topic}
- 생성일: {today}
- 작성자: 자동 생성

1. 사업 개요
{topic}에 대한 신사업 제안입니다. 본 보고서는 GPT 기반 심층 리서치를 통해 도출된 정성·정량적 인사이트를 기반으로 작성되었습니다.

2. 산업 및 시장 분석
{extract_section(research_text, ['산업 구조', '시장 규모', '트렌드', '밸류체인'])}

3. 경쟁사 벤치마킹
{extract_section(research_text, ['대표 기업', '벤치마킹', '기업별 분석'])}

4. 고객 및 수요 분석
{extract_section(research_text, ['소비자', '고객', '문화', '수요'])}

5. 수익 모델 및 BEP 분석
{extract_section(research_text, ['단가', '마진', '시장 규모', '수익', 'BEP', '시나리오'])}

6. 전략 제안 및 실행 방안
{extract_section(research_text, ['전략', '실행', '기회', '변화'])}

7. 결론 및 기대효과
{extract_section(research_text, ['기대효과', '요약', '종합'])}
"""
    return report


# 🔍 텍스트에서 키워드 기반 추출
def extract_section(text: str, keywords: list) -> str:
    lines = text.split('\n')
    extracted = []
    for line in lines:
        if any(keyword in line for keyword in keywords):
            extracted.append(line.strip())
    return '\n'.join(extracted[:10])  # 각 섹션별 최대 10줄


# 📄 워드 저장: 보고서
def save_report_as_docx(topic: str, report_text: str, filename: str = None):
    if not filename:
        filename = f"{topic}_보고서.docx"
    doc = Document()
    for line in report_text.strip().split('\n'):
        if line.strip().startswith('[') and line.strip().endswith(']'):
            doc.add_heading(line.strip().replace('[', '').replace(']', ''), level=1)
        elif line.strip().startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip().isdigit() or line.strip().startswith(tuple("123456789")):
            doc.add_heading(line.strip(), level=2)
        else:
            para = doc.add_paragraph(line.strip())
            para.style.font.size = Pt(11)
    doc.save(filename)
    print(f"📁 보고서 워드 저장 완료: {filename}")


# 📄 워드 저장: 리서치 원문
def save_raw_text_as_docx(raw_text: str, filename: str):
    doc = Document()
    doc.add_heading("GPT 심층 리서치 원문", level=1)
    for line in raw_text.strip().split('\n'):
        doc.add_paragraph(line.strip())
    doc.save(filename)
    print(f"📁 리서치 원문 워드 저장 완료: {filename}")


# 📦 전체 자동화 실행 함수 (streamlit이 부름)
def generate_all_documents(topic: str):
    today = datetime.today().strftime("%Y%m%d")
    folder = f"{topic}_{today}"
    os.makedirs(folder, exist_ok=True)

    # 1️⃣ GPT 리서치
    result = run_deep_research(topic)
    raw_text = result["raw_text"]

    # 2️⃣ 리서치 원문 저장
    docx_raw_path = os.path.join(folder, "리서치_원본.docx")
    save_raw_text_as_docx(raw_text, docx_raw_path)

    # 3️⃣ 보고서 저장
    report_text = generate_business_report(topic, raw_text)
    docx_report_path = os.path.join(folder, f"{topic}_보고서.docx")
    save_report_as_docx(topic, report_text, docx_report_path)

    # 4️⃣ PPT 저장
    slides = generate_slide_outline(topic, report_text)
    pptx_path = os.path.join(folder, f"{topic}_슬라이드.pptx")
    create_ppt_from_outline(topic, slides, pptx_path)

    # 5️⃣ ZIP 압축
    zip_path = f"{folder}.zip"
    with zipfile.ZipFile(zip_path, 'w') as zf:
        zf.write(docx_raw_path, arcname=os.path.basename(docx_raw_path))
        zf.write(docx_report_path, arcname=os.path.basename(docx_report_path))
        zf.write(pptx_path, arcname=os.path.basename(pptx_path))

    print(f"✅ ZIP 파일 생성 완료: {zip_path}")
    return zip_path, os.path.basename(zip_path)
