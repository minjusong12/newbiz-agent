import os
from datetime import datetime
from research import run_deep_research
from generate_report import generate_business_report, save_report_as_docx
from generate_slides import generate_slide_outline
from generate_pptx import create_ppt_from_outline
from docx import Document
import zipfile

# 리서치 원본을 워드로 저장
def save_raw_text_as_docx(raw_text: str, filename: str):
    doc = Document()
    doc.add_heading("GPT 심층 리서치 원문", level=1)
    for line in raw_text.strip().split('\n'):
        doc.add_paragraph(line.strip())
    doc.save(filename)
    print(f"📄 리서치 원본 워드 저장 완료: {filename}")

# 전체 자동화 파이프라인
def run_full_pipeline(topic: str, use_mock=False):
    today = datetime.today().strftime("%Y%m%d")
    folder = f"{topic}_{today}"
    os.makedirs(folder, exist_ok=True)

    # 1️⃣ GPT 리서치
    if use_mock:
        raw_text = '''
        - 산업 구조: 영세 업체와 대형 상조회 병존
        - 시장 규모: 약 1.7조 엔, 연 130만 건
        - 대표 기업: 작은 장례식, 민레비, 카마쿠라신서
        - 단가: 소개형 30만 원, 정액형 100만 원 마진
        - 소비자: 고령자 중심, 가족장 확산
        - 기대효과: 신뢰 기반 초기 확보 → 점유율 확대
        '''
    else:
        result = run_deep_research(topic)
        raw_text = result["raw_text"]

    # 2️⃣ 리서치 원문 워드 저장
    docx_raw_path = os.path.join(folder, "리서치_원본.docx")
    save_raw_text_as_docx(raw_text, docx_raw_path)

    # 3️⃣ 보고서 생성 및 저장
    report_text = generate_business_report(topic, raw_text)
    docx_report_path = os.path.join(folder, f"{topic}_보고서.docx")
    save_report_as_docx(topic, report_text, docx_report_path)

    # 4️⃣ PPT 슬라이드 생성
    slides = generate_slide_outline(topic, report_text)
    pptx_path = os.path.join(folder, f"{topic}_슬라이드.pptx")
    create_ppt_from_outline(topic, slides, pptx_path)

    # 5️⃣ 압축 파일로 패키징
    zip_path = f"{folder}.zip"
    with zipfile.ZipFile(zip_path, 'w') as zf:
        zf.write(docx_raw_path, arcname=os.path.basename(docx_raw_path))
        zf.write(docx_report_path, arcname=os.path.basename(docx_report_path))
        zf.write(pptx_path, arcname=os.path.basename(pptx_path))

    print(f"\n🎉 완료! 생성된 파일 ZIP: {zip_path}")


# 🧪 실행
if __name__ == "__main__":
    topic = input("📝 주제를 입력하세요: ")
    use_mock = input("⚙️ 외부망 차단 상태인가요? (y/n): ").strip().lower() == 'y'
    run_full_pipeline(topic, use_mock)
